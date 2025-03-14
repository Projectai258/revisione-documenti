import streamlit as st
import openai
import os
import re
import logging
import io
from dotenv import load_dotenv
from bs4 import BeautifulSoup
import markdown
from docx import Document
from PyPDF2 import PdfReader
from fpdf import FPDF

########################################
# 1) Carica variabili d'ambiente (solo Python)
########################################
load_dotenv()

########################################
# 2) PRIMO comando Streamlit
########################################
st.set_page_config(page_title="Revisione Documenti", layout="wide")

########################################
# 3) Configurazione logging e altre impostazioni Python
########################################
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

########################################
# 4) Recupera la chiave API
########################################
API_KEY = os.getenv("OPENROUTER_API_KEY") or st.secrets.get("OPENROUTER_API_KEY")
if not API_KEY:
    st.error("⚠️ Errore: API Key di OpenRouter non trovata! Impostala come variabile d'ambiente o in st.secrets.")
    st.stop()

# Inizializza il client OpenAI per OpenRouter
client = openai.OpenAI(api_key=API_KEY, base_url="https://openrouter.ai/api/v1")

########################################
# Definizione dei pattern critici
########################################
CRITICAL_PATTERNS = [
    r"\bIlias Contreas\b",
    r"\bIlias\b",
    r"\bContreas\b",
    r"\bJoey\b",
    r"\bMya\b",
    r"\bmia moglie\b",
    r"\bmia figlia\b",
    r"\bShake Your English\b",
    r"\bBarman PR\b",
    r"\bStairs Club\b",
    r"\bsocio\b",
    r"\surf\b",
    r"\sottoscritto\b",
    r"Mi chiamo .*? la mia esperienza personale\.",
    r"\bfiglio di papà\b",
    r"\bhappy our\b",
]
compiled_patterns = [re.compile(p, re.IGNORECASE) for p in CRITICAL_PATTERNS]

########################################
# Opzioni di tono per la riscrittura
########################################
TONE_OPTIONS = {
    "Stile originale": "Mantieni lo stesso stile del testo originale, stessa struttura della frase.",
    "Formale": "Riscrivi in modo formale e professionale.",
    "Informale": "Riscrivi in modo amichevole e colloquiale rivolto ad un lettore giovane.",
    "Tecnico": "Riscrivi con linguaggio tecnico e preciso.",
    "Narrativo": "Riscrivi in stile descrittivo e coinvolgente direzionato ad un pubblico giovanile.",
    "Pubblicitario": "Riscrivi in modo persuasivo, come una pubblicità.",
    "Giornalistico": "Riscrivi in tono chiaro e informativo.",
}

########################################
# Funzione per conversione da prima persona singolare a plurale
########################################
def convert_first_singular_to_plural(text):
    # Sostituzioni di esempio (puoi raffinarle ulteriormente)
    text = re.sub(r'\b[Ii]o\b', 'noi', text)
    text = re.sub(r'\b[Mm]io\b', 'nostro', text)
    text = re.sub(r'\b[Mm]ia\b', 'nostra', text)
    text = re.sub(r'\b[Mm]iei\b', 'nostri', text)
    text = re.sub(r'\b[Mm]ie\b', 'nostre', text)
    text = re.sub(r'\b[Mm]i\b', 'ci', text)
    return text

########################################
# Funzioni di supporto comuni
########################################
def extract_context(blocks, selected_block):
    """Estrae il blocco precedente e successivo per fornire contesto al modello."""
    try:
        index = blocks.index(selected_block)
    except ValueError:
        logger.error("Il blocco selezionato non è presente nella lista.")
        return "", ""
    prev_block = blocks[index - 1] if index > 0 else ""
    next_block = blocks[index + 1] if index < len(blocks) - 1 else ""
    return prev_block, next_block

def ai_rewrite_text(text, prev_text, next_text, tone):
    """Richiede all'API di riscrivere il testo in base al tono selezionato."""
    prompt = (
        f"Contesto:\nPrecedente: {prev_text}\nTesto: {text}\nSuccessivo: {next_text}\n\n"
        f"Riscrivi il 'Testo' in tono '{tone}'. Rimuovi eventuali dettagli personali o identificabili. "
        "Rispondi con UNA sola frase, senza ulteriori commenti."
    )
    try:
        response = client.chat.completions.create(
            model="google/gemini-2.0-pro-exp-02-05:free",
            messages=[{"role": "system", "content": prompt}],
            max_tokens=50
        )
        if response and hasattr(response, "choices") and response.choices:
            return response.choices[0].message.content.strip()
        error_message = "⚠️ Errore: Nessun testo valido restituito dall'API."
        logger.error(error_message)
        return error_message
    except Exception as e:
        error_message = f"⚠️ Errore nell'elaborazione: {e}"
        logger.error(error_message)
        return error_message

def process_html_content(html_content, modifications, highlight=False):
    """Sostituisce i blocchi modificati all'interno del contenuto HTML."""
    soup = BeautifulSoup(html_content, "html.parser")
    for tag in soup.find_all(["p", "span", "div", "li", "a", "h5"]):
        if tag.string:
            original = tag.string.strip()
            if original in modifications:
                mod_text = modifications[original]
                if highlight:
                    new_tag = soup.new_tag("span", style="background-color: yellow; font-weight: bold;")
                    new_tag.string = mod_text
                    tag.string.replace_with("")
                    tag.append(new_tag)
                else:
                    tag.string.replace_with(mod_text)
    return str(soup)

def generate_html_preview(blocks, modifications, highlight=False):
    """Genera un'anteprima HTML evidenziata."""
    html = ""
    for block in blocks:
        mod_text = modifications.get(block, block)
        if highlight:
            html += f'<p><span style="background-color: yellow; font-weight: bold;">{mod_text}</span></p>'
        else:
            html += f"<p>{mod_text}</p>"
    return html

def process_file_content(file_content, file_extension):
    """Elabora il contenuto per file HTML/Markdown e ritorna (lista_blocchi, contenuto_html)."""
    if file_extension == "html":
        html_content = file_content
    elif file_extension == "md":
        html_content = markdown.markdown(file_content)
    else:
        html_content = ""
    if html_content:
        soup = BeautifulSoup(html_content, "html.parser")
        blocks = [tag.string.strip() for tag in soup.find_all(["p", "span", "div", "li", "a", "h5"]) if tag.string]
        return blocks, html_content
    return [], ""

def process_doc_file(uploaded_file):
    """Estrae i paragrafi da un file Word."""
    try:
        doc = Document(uploaded_file)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return paragraphs
    except Exception as e:
        st.error(f"Errore nell'apertura del file Word: {e}")
        st.stop()

def process_pdf_file(uploaded_file):
    """Estrae il testo da un file PDF (usato per mostrare i blocchi da revisionare)."""
    try:
        pdf_reader = PdfReader(uploaded_file)
    except Exception as e:
        st.error(f"Errore nell'apertura del file PDF: {e}")
        st.stop()
    paragraphs = []
    for page in pdf_reader.pages:
        text = page.extract_text()
        if text:
            paragraphs.extend([line.strip() for line in text.split("\n") if line.strip()])
    return paragraphs

def filtra_blocchi(blocchi):
    """Filtra i blocchi che corrispondono ai pattern critici."""
    return {f"{i}_{b}": b for i, b in enumerate(blocchi) if any(pattern.search(b) for pattern in compiled_patterns)}

########################################
# Funzione per elaborare PDF con overlay (usando PyMuPDF)
########################################
def process_pdf_with_overlay(uploaded_file, modifications):
    """
    Apre il PDF originale con PyMuPDF (fitz), cerca i blocchi di testo che contengono il testo originale (presenti in modifications),
    aggiunge un'annotazione di redazione per cancellare il testo originale e inserisce il testo revisionato nello stesso rettangolo.
    Ritorna il PDF modificato come bytes.
    """
    import fitz  # PyMuPDF
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    for page in doc:
        blocks = page.get_text("blocks")
        for b in blocks:
            block_text = b[4].strip()
            for original, revised in modifications.items():
                if original in block_text:
                    rect = fitz.Rect(b[0], b[1], b[2], b[3])
                    page.add_redact_annot(rect, fill=(1,1,1))
                    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
                    # Inserisce il testo revisionato nello stesso rettangolo, centrato
                    page.insert_textbox(rect, revised, fontsize=12, fontname="helv", align=1)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

########################################
# Selezione modalità e flag di conversione globale
########################################
# Modalità di revisione (rimane la scelta per blocchi)
modalita = st.radio(
    "Modalità di revisione:",
    ("Revisiona blocchi corrispondenti", "Conversione intera (solo)", "Revisiona blocchi e applica conversione globale")
)
# Se la modalità scelta è "Revisiona blocchi e applica conversione globale", il flag sarà True
global_conversion = modalita == "Revisiona blocchi e applica conversione globale"

########################################
# Logica principale Streamlit
########################################

st.title("📄 Revisione Documenti")
st.write("Carica un file (HTML, Markdown, Word o PDF) e scegli come intervenire sul testo.")

uploaded_file = st.file_uploader("📂 Seleziona un file (html, md, doc, docx, pdf)", type=["html", "md", "doc", "docx", "pdf"])

if uploaded_file is not None:
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    # Modalità "Conversione intera (solo)"
    if modalita == "Conversione intera (solo)":
        if file_extension in ["html", "md"]:
            file_content = uploaded_file.read().decode("utf-8")
            converted_text = convert_first_singular_to_plural(file_content)
            st.subheader("📌 Testo Revisionato (Conversione Intera)")
            st.code(converted_text, language="html" if file_extension=="html" else "plaintext")
            if st.button("📥 Scarica File Revisionato"):
                st.download_button("Scarica Revisionato", converted_text.encode("utf-8"), "document_revised.html" if file_extension=="html" else "document_revised.txt", "text/html" if file_extension=="html" else "text/plain")
        elif file_extension in ["doc", "docx"]:
            paragraphs = process_doc_file(uploaded_file)
            full_text = "\n".join(paragraphs)
            converted_text = convert_first_singular_to_plural(full_text)
            st.subheader("📌 Testo Revisionato (Conversione Intera)")
            st.code(converted_text, language="plaintext")
            if st.button("📥 Scarica Documento Revisionato"):
                new_doc = Document()
                new_doc.add_paragraph(converted_text)
                buffer = io.BytesIO()
                new_doc.save(buffer)
                st.download_button("Scarica Documento Revisionato", buffer.getvalue(), "document_revised.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        elif file_extension == "pdf":
            paragraphs = process_pdf_file(uploaded_file)
            full_text = "\n".join(paragraphs)
            converted_text = convert_first_singular_to_plural(full_text)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, converted_text)
            buffer = io.BytesIO()
            pdf.output(buffer, 'F')
            st.subheader("📌 PDF Revisionato (Conversione Intera)")
            if st.button("📥 Scarica PDF Revisionato"):
                st.download_button("Scarica PDF Revisionato", buffer.getvalue(), "document_revised.pdf", "application/pdf")
    
    # Modalità "Revisiona blocchi corrispondenti" oppure "Revisiona blocchi e applica conversione globale"
    else:
        modifications = {}
        if file_extension in ["html", "md"]:
            file_content = uploaded_file.read().decode("utf-8")
            blocchi, html_content = process_file_content(file_content, file_extension)
            blocchi_da_revisionare = filtra_blocchi(blocchi)
            if blocchi_da_revisionare:
                st.subheader("📌 Blocchi da revisionare")
                progress_text = st.empty()
                progress_bar = st.progress(0)
                total = len(blocchi_da_revisionare)
                count = 0
                for uid, blocco in blocchi_da_revisionare.items():
                    st.markdown(f"**{blocco}**")
                    azione = st.radio("Azione per questo blocco:", ["Riscrivi", "Elimina", "Ignora"], key=f"action_{uid}")
                    if azione == "Riscrivi":
                        tono = st.selectbox("Scegli il tono:", list(TONE_OPTIONS.keys()), key=f"tone_{uid}")
                        prev_blocco, next_blocco = extract_context(blocchi, blocco)
                        attempts = 0
                        mod_blocco = ""
                        while attempts < 3:
                            mod_blocco = ai_rewrite_text(blocco, prev_blocco, next_blocco, tono)
                            if "Errore" not in mod_blocco:
                                break
                            attempts += 1
                        modifications[blocco] = mod_blocco
                    elif azione == "Elimina":
                        modifications[blocco] = ""
                    else:
                        modifications[blocco] = blocco
                    count += 1
                    progress_bar.progress(count / total)
                    progress_text.text(f"Elaborati {count} di {total} blocchi...")
                # Applica conversione globale al testo finale se il flag è attivo
                final_content = process_html_content(html_content, modifications, highlight=True)
                if global_conversion:
                    final_content = convert_first_singular_to_plural(final_content)
                if st.button("✍️ Genera Documento Revisionato"):
                    with st.spinner("🔄 Riscrittura in corso..."):
                        st.session_state["preview_content"] = final_content
                    st.success("✅ Revisione completata!")
                    st.subheader("🌍 Anteprima con Testo Revisionato")
                    st.components.v1.html(st.session_state["preview_content"], height=500, scrolling=True)
                    st.download_button("📥 Scarica HTML Revisionato", st.session_state["preview_content"].encode("utf-8"), "document_revised.html", "text/html")
            else:
                st.info("Non sono state trovate corrispondenze per i criteri di ricerca nel testo.")
        
        elif file_extension in ["doc", "docx"]:
            paragrafi = process_doc_file(uploaded_file)
            blocchi_da_revisionare = filtra_blocchi(paragrafi)
            if blocchi_da_revisionare:
                st.subheader("📌 Paragrafi da revisionare")
                progress_text = st.empty()
                progress_bar = st.progress(0)
                total = len(blocchi_da_revisionare)
                count = 0
                for uid, paragrafo in blocchi_da_revisionare.items():
                    st.markdown(f"**{paragrafo}**")
                    azione = st.radio("Azione per questo paragrafo:", ["Riscrivi", "Elimina", "Ignora"], key=f"action_{uid}")
                    if azione == "Riscrivi":
                        tono = st.selectbox("Scegli il tono:", list(TONE_OPTIONS.keys()), key=f"tone_{uid}")
                        prev_par, next_par = extract_context(paragrafi, paragrafo)
                        attempts = 0
                        mod_par = ""
                        while attempts < 3:
                            mod_par = ai_rewrite_text(paragrafo, prev_par, next_par, tono)
                            if "Errore" not in mod_par:
                                break
                            attempts += 1
                        modifications[paragrafo] = mod_par
                    elif azione == "Elimina":
                        modifications[paragrafo] = ""
                    else:
                        modifications[paragrafo] = paragrafo
                    count += 1
                    progress_bar.progress(count / total)
                    progress_text.text(f"Elaborati {count} di {total} paragrafi...")
                # Crea il documento Word e applica conversione globale se richiesto
                full_text = "\n".join([modifications.get(p, p) for p in paragrafi])
                if global_conversion:
                    full_text = convert_first_singular_to_plural(full_text)
                new_doc = Document()
                new_doc.add_paragraph(full_text)
                buffer = io.BytesIO()
                new_doc.save(buffer)
                st.success("✅ Revisione completata!")
                st.subheader("🌍 Anteprima Testo (Word)")
                st.download_button("📥 Scarica Documento Word Revisionato", buffer.getvalue(), "document_revised.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.info("Non sono state trovate corrispondenze per i criteri di ricerca nel documento Word.")
        
        elif file_extension == "pdf":
            paragrafi = process_pdf_file(uploaded_file)
            blocchi_da_revisionare = filtra_blocchi(paragrafi)
            if blocchi_da_revisionare:
                st.subheader("📌 Blocchi di testo da revisionare (PDF)")
                progress_text = st.empty()
                progress_bar = st.progress(0)
                total = len(blocchi_da_revisionare)
                count = 0
                for uid, blocco in blocchi_da_revisionare.items():
                    st.markdown(f"**{blocco}**")
                    azione = st.radio("Azione per questo blocco:", ["Riscrivi", "Elimina", "Ignora"], key=f"action_{uid}")
                    if azione == "Riscrivi":
                        tono = st.selectbox("Scegli il tono:", list(TONE_OPTIONS.keys()), key=f"tone_{uid}")
                        prev_blocco, next_blocco = extract_context(paragrafi, blocco)
                        attempts = 0
                        mod_blocco = ""
                        while attempts < 3:
                            mod_blocco = ai_rewrite_text(blocco, prev_blocco, next_blocco, tono)
                            if "Errore" not in mod_blocco:
                                break
                            attempts += 1
                        modifications[blocco] = mod_blocco
                    elif azione == "Elimina":
                        modifications[blocco] = ""
                    else:
                        modifications[blocco] = blocco
                    count += 1
                    progress_bar.progress(count / total)
                    progress_text.text(f"Elaborati {count} di {total} blocchi...")
                # Se il flag di conversione globale è attivo, applica la conversione anche ai testi di modifica
                if global_conversion:
                    for key in modifications:
                        modifications[key] = convert_first_singular_to_plural(modifications[key])
                if st.button("✍️ Genera PDF Revisionato"):
                    with st.spinner("🔄 Riscrittura in corso..."):
                        revised_pdf = process_pdf_with_overlay(uploaded_file, modifications)
                    st.success("✅ Revisione completata!")
                    st.download_button("📥 Scarica PDF Revisionato", revised_pdf, "document_revised.pdf", "application/pdf")
            else:
                st.info("Non sono state trovate corrispondenze per i criteri di ricerca nel documento PDF.")
